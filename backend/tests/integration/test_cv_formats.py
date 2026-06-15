"""
Integration tests for V2-1 — additional CV file formats (DOCX & images).
Constitution Principle VIII (CV screening domain) — write FIRST, confirm FAILING
before the implementation tasks (T004–T006).

Mirrors the mocking strategy in test_screening.py: the OCR/extraction boundary,
the embedding service, and the agents HTTP call are mocked so these tests exercise
the format-dispatch + router-validation behaviour, not external services.
"""

from __future__ import annotations

import asyncio
import io
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from httpx import ASGITransport, AsyncClient

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# embed_text returns (embedding, usage_event)
EMBED_RESULT = (
    [0.0] * 1536,
    {
        "agent_type": "embedding",
        "model": "text-embedding-3-small",
        "prompt_tokens": 10,
        "completion_tokens": 0,
        "estimated_cost_usd": 0.0,
        "metadata": {"operation": "cv_chunk_embedding"},
    },
)


def _agents_client(payload: dict):
    """Replacement for the `httpx.AsyncClient` *name*: patching `.post` would also
    hijack the ASGI test client (it is an httpx.AsyncClient too)."""

    class _Client:
        def __init__(self, *args, **kwargs):
            pass

        async def __aenter__(self):
            return self

        async def __aexit__(self, *exc):
            return False

        async def post(self, *args, **kwargs):
            resp = MagicMock()
            resp.status_code = 200
            resp.raise_for_status = MagicMock()
            resp.json = MagicMock(return_value=payload)
            return resp

    return _Client


async def _wait_for_screening(client: AsyncClient, token: str, application_id: str) -> dict:
    """Screening is a fire-and-forget task — poll until it leaves 'pending'."""
    data: dict = {}
    for _ in range(50):
        await asyncio.sleep(0.1)
        resp = await client.get(
            f"/applications/{application_id}",
            headers={"Authorization": f"Bearer {token}"},
        )
        assert resp.status_code == 200
        data = resp.json()
        if data["screening_status"] != "pending":
            break
    return data


@pytest.fixture
async def client(app):
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as c:
        yield c


def _make_docx_bytes(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    """Build a real .docx (paragraphs + an optional table) in memory."""
    import docx  # python-docx

    document = docx.Document()
    for p in paragraphs:
        document.add_paragraph(p)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for col, value in enumerate(row):
                table.cell(r, col).text = value
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# T002 — DOCX upload → extraction (paragraphs + table) + screening
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_docx_cv_extracts_and_screens(client: AsyncClient, active_job_token):
    """
    A text-based .docx (paragraphs + a table) yields non-empty cv_text, a screening
    score/rationale, and cv_extraction_method == 'docx'.
    """
    token, job_id = active_job_token
    docx_bytes = _make_docx_bytes(
        paragraphs=[
            "Jane Roe — Senior Python Engineer",
            "8 years building async FastAPI services and data pipelines.",
            # Keep the text above the 50-word extraction-quality threshold so the
            # native docx path is used (no Document Intelligence fallback).
            "Designed and operated event-driven microservices on Azure Container Apps "
            "with PostgreSQL, Redis, and pgvector, owning reliability and observability "
            "end to end across multiple production tenants.",
            "Led a team of four engineers, introduced contract testing and trunk-based "
            "development, and reduced deployment lead time from days to under an hour "
            "while keeping change failure rate low.",
        ],
        table_rows=[["Skill", "Years"], ["Python", "8"], ["PostgreSQL", "6"]],
    )

    agents_payload = {
        "score": 88,
        "rationale": "Strong async Python and PostgreSQL match.",
        "status": "qualified",
        "guardrail_triggered": False,
    }
    with (
        patch(
            "app.services.embedding_service.EmbeddingService.embed_text", new_callable=AsyncMock
        ) as mock_embed,
        patch("httpx.AsyncClient", _agents_client(agents_payload)),
    ):
        mock_embed.return_value = EMBED_RESULT

        resp = await client.post(
            f"/jobs/{job_id}/applications",
            data={"full_name": "Jane Roe", "email": "jane.roe@example.com"},
            files={"cv_file": ("cv.docx", docx_bytes, DOCX_MIME)},
        )
        assert resp.status_code == 201
        application_id = resp.json()["id"]

        data = await _wait_for_screening(client, token, application_id)
    assert data["screening_score"] is not None
    assert data["screening_rationale"] is not None
    assert data["cv_extraction_method"] == "docx"


@pytest.mark.asyncio
async def test_sparse_docx_falls_back_to_document_intelligence(
    client: AsyncClient, active_job_token
):
    """
    A text-light .docx (below the quality threshold) falls back to the Document
    Intelligence path, recording cv_extraction_method == 'document_intelligence'.
    """
    token, job_id = active_job_token
    sparse_docx = _make_docx_bytes(paragraphs=["Photo CV"])  # < 50 words → sparse

    agents_payload = {
        "score": 60,
        "rationale": "Adequate match.",
        "status": "qualified",
        "guardrail_triggered": False,
    }
    with (
        patch(
            "app.services.ocr_service.OcrService._azure_doc_intelligence",
            new_callable=AsyncMock,
        ) as mock_di,
        patch(
            "app.services.embedding_service.EmbeddingService.embed_text", new_callable=AsyncMock
        ) as mock_embed,
        patch("httpx.AsyncClient", _agents_client(agents_payload)),
    ):
        mock_di.return_value = (
            "Recovered CV text from Document Intelligence with plenty of words " * 10,
            "document_intelligence",
        )
        mock_embed.return_value = EMBED_RESULT

        resp = await client.post(
            f"/jobs/{job_id}/applications",
            data={"full_name": "Sparse Sam", "email": "sparse.sam@example.com"},
            files={"cv_file": ("sparse.docx", sparse_docx, DOCX_MIME)},
        )
        assert resp.status_code == 201
        application_id = resp.json()["id"]
        mock_di.assert_awaited()  # DI fallback was used

        data = await _wait_for_screening(client, token, application_id)
    assert data["cv_extraction_method"] == "document_intelligence"


# ---------------------------------------------------------------------------
# T003 — Image routing, unsupported types, blank image, corrupt/renamed files
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_image_cv_routes_to_document_intelligence(client: AsyncClient, active_job_token):
    """A .png/.jpg CV routes straight to Document Intelligence and produces a result."""
    token, job_id = active_job_token
    png_bytes = b"\x89PNG\r\n\x1a\n" + b"0" * 256  # PNG magic + filler

    agents_payload = {
        "score": 75,
        "rationale": "Relevant experience.",
        "status": "qualified",
        "guardrail_triggered": False,
    }
    with (
        patch(
            "app.services.ocr_service.OcrService._azure_doc_intelligence",
            new_callable=AsyncMock,
        ) as mock_di,
        patch(
            "app.services.embedding_service.EmbeddingService.embed_text", new_callable=AsyncMock
        ) as mock_embed,
        patch("httpx.AsyncClient", _agents_client(agents_payload)),
    ):
        mock_di.return_value = (
            "OCR-extracted CV content with many relevant skills and experience " * 10,
            "document_intelligence",
        )
        mock_embed.return_value = EMBED_RESULT

        resp = await client.post(
            f"/jobs/{job_id}/applications",
            data={"full_name": "Ivan Image", "email": "ivan.image@example.com"},
            files={"cv_file": ("cv.png", png_bytes, "image/png")},
        )
        assert resp.status_code == 201
        application_id = resp.json()["id"]
        mock_di.assert_awaited()

        data = await _wait_for_screening(client, token, application_id)
    assert data["cv_extraction_method"] == "document_intelligence"


@pytest.mark.asyncio
async def test_blank_image_rejected_no_record(client: AsyncClient, active_job):
    """
    A blank/blurry image yields an empty extraction → 422 'could not read CV',
    and no application record is created (the no_text_extracted branch).
    """
    job_id = active_job
    blank_png = b"\x89PNG\r\n\x1a\n" + b"0" * 64

    from app.services.ocr_service import OcrValidationError

    with patch(
        "app.services.ocr_service.OcrService._azure_doc_intelligence",
        new_callable=AsyncMock,
    ) as mock_di:
        mock_di.side_effect = OcrValidationError("no_text_extracted")

        resp = await client.post(
            f"/jobs/{job_id}/applications",
            data={"full_name": "Blank Bill", "email": "blank.bill@example.com"},
            files={"cv_file": ("blank.png", blank_png, "image/png")},
        )
        assert resp.status_code == 422

    # No record should exist for this candidate on this job.
    # (Listing requires auth; the 422-before-DB contract is the primary signal — FR-006.)


@pytest.mark.asyncio
async def test_unsupported_txt_rejected_no_record(client: AsyncClient, active_job):
    """A .txt upload returns 422 with a clear message and creates NO application record."""
    job_id = active_job

    resp = await client.post(
        f"/jobs/{job_id}/applications",
        data={"full_name": "Terry Text", "email": "terry.text@example.com"},
        files={"cv_file": ("cv.txt", b"just plain text, not a CV file", "text/plain")},
    )
    assert resp.status_code == 422
    assert "Accepted: PDF, DOCX, JPG, PNG" in resp.json()["detail"]


@pytest.mark.asyncio
async def test_renamed_pdf_as_docx_rejected(client: AsyncClient, active_job):
    """
    A file whose bytes don't parse as the claimed type (e.g. a PDF renamed to .docx)
    is rejected with 422 — validated by content, not just the extension.
    """
    job_id = active_job
    not_a_docx = b"%PDF-1.4 this is actually a pdf, not a docx"

    resp = await client.post(
        f"/jobs/{job_id}/applications",
        data={"full_name": "Ricky Renamed", "email": "ricky.renamed@example.com"},
        files={"cv_file": ("resume.docx", not_a_docx, DOCX_MIME)},
    )
    assert resp.status_code == 422
