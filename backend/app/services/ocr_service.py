from __future__ import annotations

import io

import anyio

# Extensions we know how to dispatch, and the MIME types we map onto them.
_SUPPORTED_EXTS = {"pdf", "docx", "jpg", "jpeg", "png"}
_CONTENT_TYPE_EXT = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "image/jpeg": "jpg",
    "image/png": "png",
}


class OcrValidationError(ValueError):
    pass


class OcrService:
    """
    Extract text from a CV file (PDF, DOCX, or image).

    Strategy, dispatched by file type:
    1. PDF   → PyMuPDF (fast, native text), then Azure Document Intelligence on sparse text.
    2. DOCX  → python-docx (paragraphs + tables), then Document Intelligence on sparse text.
    3. Image → Azure Document Intelligence directly (no local text layer to try).

    Falls back when extracted text is below the quality threshold (word count < 50 OR
    printable ratio < 0.90). Raises OcrValidationError for corrupt/encrypted/unreadable
    files or unsupported types so the caller can return a 422 with no record created.
    """

    async def extract(
        self,
        file_bytes: bytes,
        filename: str = "cv.pdf",
        content_type: str | None = None,
    ) -> tuple[str, str]:
        """
        Returns (cv_text, extraction_method).
        extraction_method: 'pymupdf' | 'docx' | 'document_intelligence'
        """
        ext = self._detect_ext(filename, content_type)

        if ext == "pdf":
            text, method = await self._try_pymupdf(file_bytes)
            if self._quality_ok(text):
                return text, method
            return await self._azure_doc_intelligence(file_bytes)

        if ext == "docx":
            text, method = await self._try_docx(file_bytes)
            if self._quality_ok(text):
                return text, method
            return await self._azure_doc_intelligence(file_bytes)

        if ext in ("jpg", "jpeg", "png"):
            # Images have no text layer — go straight to Document Intelligence.
            return await self._azure_doc_intelligence(file_bytes)

        raise OcrValidationError(f"unsupported_type: {ext or 'unknown'}")

    def _detect_ext(self, filename: str | None, content_type: str | None) -> str:
        """Resolve a normalized extension from the filename, then the declared MIME."""
        name = (filename or "").lower()
        if "." in name:
            ext = name.rsplit(".", 1)[1]
            if ext in _SUPPORTED_EXTS:
                return ext
        if content_type and content_type in _CONTENT_TYPE_EXT:
            return _CONTENT_TYPE_EXT[content_type]
        # Unknown — return whatever extension we saw (if any) so the caller errors clearly.
        return name.rsplit(".", 1)[1] if "." in name else ""

    async def _try_pymupdf(self, file_bytes: bytes) -> tuple[str, str]:
        import fitz  # PyMuPDF

        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
        except Exception as exc:
            raise OcrValidationError(f"corrupted_pdf: {exc}") from exc

        if doc.is_encrypted:
            raise OcrValidationError("encrypted_pdf")

        pages_text = []
        for page in doc:
            pages_text.append(page.get_text())
        doc.close()
        return "\n".join(pages_text), "pymupdf"

    async def _try_docx(self, file_bytes: bytes) -> tuple[str, str]:
        # python-docx is synchronous/CPU-bound — run it off the event loop so a large
        # DOCX parse does not block the async server (Constitution Principle II).
        text = await anyio.to_thread.run_sync(self._parse_docx_sync, file_bytes)
        return text, "docx"

    @staticmethod
    def _parse_docx_sync(file_bytes: bytes) -> str:
        import docx  # python-docx

        try:
            document = docx.Document(io.BytesIO(file_bytes))
        except Exception as exc:
            # Corrupt, password-protected, or not actually a .docx (e.g. renamed file).
            raise OcrValidationError(f"corrupted_docx: {exc}") from exc

        parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)

    def _quality_ok(self, text: str) -> bool:
        words = text.split()
        if len(words) < 50:
            return False
        printable = sum(1 for c in text if c.isprintable())
        total = max(len(text), 1)
        if printable / total < 0.90:
            return False
        return True

    async def _azure_doc_intelligence(self, file_bytes: bytes) -> tuple[str, str]:
        from app.config import get_settings
        from azure.ai.formrecognizer.aio import DocumentAnalysisClient
        from azure.core.credentials import AzureKeyCredential

        settings = get_settings()
        client = DocumentAnalysisClient(
            endpoint=settings.AZURE_FORM_RECOGNIZER_ENDPOINT,
            credential=AzureKeyCredential(settings.AZURE_FORM_RECOGNIZER_KEY),
        )
        async with client:
            poller = await client.begin_analyze_document("prebuilt-document", file_bytes)
            result = await poller.result()

        text = "\n".join(p.content for p in result.paragraphs or [])
        if not text.strip():
            raise OcrValidationError("no_text_extracted")
        return text, "document_intelligence"
